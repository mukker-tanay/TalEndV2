"""
Standalone QC harness for CV NAME extraction accuracy.

This script is fully self-contained. It does NOT import or modify any app code
(app/utils/parser.py, gemini_parser.py, etc.) — it re-implements text extraction
and the Gemini prompt locally so it can be copied next to a folder of CVs and run
anywhere. Nothing here touches the production codebase or the database.

Goal: measure, on YOUR real (often messy, unstructured Indian) CVs, how accurate
each name-extraction method is — so you can decide how much you still need Gemini.

--------------------------------------------------------------------------------
USAGE
--------------------------------------------------------------------------------
1) Label a folder of CVs (creates ground truth; CVs are unlabeled):

    python qc_name_eval.py label --dir /path/to/cvs --out labels.csv
    # add --seed filename   to pre-fill a guess you can accept with Enter

2) Evaluate every method against the labels:

    python qc_name_eval.py eval --dir /path/to/cvs --labels labels.csv --out results.csv
    # add --no-gemini       to skip the Gemini method (e.g. when out of quota)

Reads GEMINI_API_KEY / GEMINI_MODEL from the environment (.env) for the gemini method.

--------------------------------------------------------------------------------
METHODS UNDER TEST
--------------------------------------------------------------------------------
  filename   - parse the file name (Rahul_Sharma_Resume.pdf -> Rahul Sharma)
  topline    - first plausible name line near the top of the CV
  email      - derive from the email local-part (rahul.sharma@ -> Rahul Sharma)
  spacy_ner  - first PERSON entity from spaCy en_core_web_sm
  ensemble   - combine the above, validated/scored with indian-namematch
  gemini     - the production approach (accuracy ceiling / baseline)

Comparison uses indian-namematch (soundex + Indian spelling variants + title
stripping) when available, plus exact / token-set normalization fallbacks.
"""

import argparse
import csv
import os
import re
import sys

# ----------------------------------------------------------------------------
# Optional dependencies — degrade gracefully, report what's available.
# ----------------------------------------------------------------------------
try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

from functools import lru_cache

try:
    from indian_namematch import fuzzymatch

    @lru_cache(maxsize=200000)
    def inm_match(a, b):
        try:
            return str(fuzzymatch.single_compare(a, b)).strip().lower() == "match"
        except Exception:
            return False
    HAS_INM = True
except Exception:
    def inm_match(a, b):
        return False
    HAS_INM = False

_SPACY_NLP = None
def _get_nlp():
    global _SPACY_NLP
    if _SPACY_NLP is None:
        import spacy
        _SPACY_NLP = spacy.load("en_core_web_sm")
    return _SPACY_NLP


# ----------------------------------------------------------------------------
# Text extraction (mirrors production exactly: PyMuPDF get_text("text"))
# ----------------------------------------------------------------------------
# Which PDF text extractor to use: "auto" | "fitz" | "pdfminer" | "pypdfium2".
# Production uses "fitz" (PyMuPDF); QC found "pdfminer" recovers the name far better.
PDF_ENGINE = "auto"

# spaCy NER candidate in the ML pool: weak signal (~25% alone) and its repeated
# nlp() calls balloon RAM to multiple GB over a full dataset. Off by default.
USE_SPACY = False


def extract_text(path):
    ext = path.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return _pdf_text(path)
    if ext == "docx":
        return _docx_text(path)
    return ""


def _pdf_fitz(path):
    import fitz  # PyMuPDF, same as production
    doc = fitz.open(path)
    text = "".join((pg.get_text("text") or "") for pg in doc)
    doc.close()
    return text.strip()


def _pdf_pdfminer(path):
    from pdfminer.high_level import extract_text as _pm
    return (_pm(path) or "").strip()


def _pdf_pypdfium2(path):
    import pypdfium2 as pdfium
    pdf = pdfium.PdfDocument(path)
    return "\n".join(pg.get_textpage().get_text_range() for pg in pdf).strip()


def _pdf_text(path):
    engines = {"fitz": _pdf_fitz, "pdfminer": _pdf_pdfminer, "pypdfium2": _pdf_pypdfium2}
    order = [PDF_ENGINE] if PDF_ENGINE in engines else ["fitz", "pdfminer", "pypdfium2"]
    if PDF_ENGINE == "auto":
        order = ["fitz", "pdfminer", "pypdfium2"]
    last_err = None
    for name in order:
        try:
            return engines[name](path)
        except ImportError:
            continue
        except Exception as e:
            last_err = e
    print(f"  [warn] cannot read PDF {os.path.basename(path)}: {last_err}")
    return ""


def _docx_text(path):
    try:
        import docx
        d = docx.Document(path)
        return "\n".join(p.text for p in d.paragraphs).strip()
    except Exception as e:
        print(f"  [warn] cannot read DOCX {os.path.basename(path)}: {e}")
        return ""


# ----------------------------------------------------------------------------
# Name normalization & comparison
# ----------------------------------------------------------------------------
_TITLES = {"mr", "mrs", "ms", "miss", "dr", "prof", "shri", "smt", "kum", "er"}
_FILENAME_NOISE = {
    "resume", "cv", "curriculum", "vitae", "profile", "final", "updated",
    "new", "latest", "copy", "doc", "document",
    # job-portal prefixes
    "naukri", "linkedin", "iimjobs", "monster", "shine", "indeed", "foundit",
    "hirist", "instahyre", "cutshort", "updazz", "timesjobs",
}


def normalize(name):
    if not name:
        return ""
    name = re.sub(r"[^A-Za-z\s]", " ", name)
    tokens = [t for t in name.lower().split() if t and t not in _TITLES]
    return " ".join(tokens).strip()


def is_match(pred, truth):
    """True if predicted name matches ground truth under any tier."""
    if not pred or not truth:
        return False
    np_, nt = normalize(pred), normalize(truth)
    if not np_ or not nt:
        return False
    if np_ == nt:
        return True
    if sorted(np_.split()) == sorted(nt.split()):   # order-insensitive
        return True
    return inm_match(pred, truth)                    # phonetic / Indian variants


# ----------------------------------------------------------------------------
# Extraction methods: (text, filename) -> predicted name or None
# ----------------------------------------------------------------------------
def m_filename(text, filename):
    base = os.path.splitext(os.path.basename(filename))[0]
    base = re.sub(r"[_\-\.]+", " ", base)
    base = re.sub(r"\d+", " ", base)
    tokens = [t for t in base.split() if t.lower() not in _FILENAME_NOISE and len(t) > 1]
    tokens = [t for t in tokens if t.isalpha()]
    if 1 <= len(tokens) <= 4:
        return " ".join(w.capitalize() for w in tokens)
    return None


def m_email(text, filename):
    m = re.search(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", text)
    if not m:
        return None
    local = m.group(0).split("@")[0]
    parts = [p for p in re.split(r"[._\-0-9]+", local) if p.isalpha() and len(p) > 1]
    if 1 <= len(parts) <= 3:
        return " ".join(p.capitalize() for p in parts)
    return None


_HEADER_WORDS = {"resume", "curriculum", "vitae", "cv", "profile", "name", "contact"}
def m_topline(text, filename):
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    for ln in lines[:12]:
        low = ln.lower()
        if "@" in ln or "http" in low or re.search(r"\d", ln):
            continue
        if any(w in low for w in _HEADER_WORDS):
            continue
        tokens = ln.split()
        if 2 <= len(tokens) <= 4 and all(re.fullmatch(r"[A-Za-z.]+", t) for t in tokens):
            return " ".join(t.capitalize().rstrip(".") for t in tokens)
    return None


def m_spacy(text, filename):
    try:
        nlp = _get_nlp()
    except Exception:
        return None
    doc = nlp(text[:500])
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text.strip()
    return None


def m_ensemble(text, filename):
    """Collect candidates; prefer ones that agree (validated by indian-namematch)."""
    cands = []
    for fn in (m_filename, m_topline, m_email, m_spacy):
        v = fn(text, filename)
        if v and normalize(v):
            cands.append(v)
    if not cands:
        return None
    # Prefer a candidate that another source agrees with.
    for i, a in enumerate(cands):
        for j, b in enumerate(cands):
            if i != j and is_match(a, b):
                return a
    # Otherwise prefer a clean 2-token alphabetic name, else first candidate.
    for c in cands:
        if len(normalize(c).split()) == 2:
            return c
    return cands[0]


_GEMINI_PROMPT = (
    "Look at the beginning of this resume and find the candidate's full name.\n"
    "Return ONLY the full name as a plain string. No explanation, no JSON.\n"
    "If you truly cannot find a name, return the single word null.\n\n"
    'Resume text:\n"""{txt}"""\n'
)


def m_gemini(text, filename):
    import json
    import requests
    key = os.getenv("GEMINI_API_KEY")
    model = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")
    if not key:
        return None
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    payload = {"contents": [{"parts": [{"text": _GEMINI_PROMPT.format(txt=text[:1500])}]}]}
    try:
        r = requests.post(f"{url}?key={key}",
                          headers={"Content-Type": "application/json"},
                          data=json.dumps(payload), timeout=30)
        if r.status_code != 200:
            print(f"  [warn] gemini HTTP {r.status_code}")
            return None
        out = r.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
        return None if out.lower() in ("null", "none", "") else out
    except Exception as e:
        print(f"  [warn] gemini error: {e}")
        return None


_CERTS = {"pmp", "cfa", "mba", "ca", "cpa", "phd", "pgdm", "be", "btech",
          "mtech", "msc", "bsc", "pg", "ug", "pmi", "csm", "scrum"}


def _camel(tok):
    return re.findall(r"[A-Z]+(?=[A-Z][a-z])|[A-Z][a-z]+|[A-Z]+|[a-z]+", tok)


def _fname_name_tokens(filename):
    base = os.path.splitext(os.path.basename(filename))[0]
    base = re.split(r"[\[\(]", base)[0]                  # drop [14y_0m] etc.
    base = re.sub(r"[_\-\.]+", " ", base).strip()
    toks = [t for t in base.split() if t.lower() not in _FILENAME_NOISE]
    out = []
    for t in toks:
        out += _camel(t)
    return [t for t in out if t.isalpha()]               # keep single-letter initials


def _email_name_tokens(text):
    m = re.search(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", text)
    if not m:
        return []
    local = m.group(0).split("@")[0]
    out = []
    for p in re.split(r"[._\-0-9]+", local):
        out += _camel(p)
    return [p.lower() for p in out if p.isalpha() and len(p) > 1]


def m_reconcile(text, filename):
    """Recover the name from the filename, using email tokens to confirm/segment.

    Handles camelCase-jammed and ALL-CAPS portal exports (DeepakPandeyPMP,
    ANIKETBISWAS) and drops cert/noise tokens the email doesn't confirm.
    """
    ftoks = _fname_name_tokens(filename)
    if not ftoks:
        return None
    etoks = set(_email_name_tokens(text))
    seg = []
    for t in ftoks:
        tl = t.lower()
        # split an ALL-CAPS blob using email tokens (ANIKETBISWAS -> aniket biswas)
        if t.isupper() and len(t) > 6 and tl not in etoks:
            rem, parts = tl, []
            for e in sorted(etoks, key=len, reverse=True):
                if rem.startswith(e):
                    parts.append(e)
                    rem = rem[len(e):]
            if not rem and parts:
                seg += parts
                continue
        seg.append(tl)
    # drop cert-like tokens unless the email confirms them
    keep = [t for t in seg if t not in _CERTS or t in etoks]
    final = [t.capitalize() for t in keep if t.isalpha()]
    return " ".join(final) if final else None


# Order matters for reporting; gemini last.
LOCAL_METHODS = [
    ("filename", m_filename),
    ("topline", m_topline),
    ("email", m_email),
    ("spacy_ner", m_spacy),
    ("ensemble", m_ensemble),
]
GEMINI_METHOD = ("gemini", m_gemini)


# ----------------------------------------------------------------------------
# ML candidate-ranker support: generate candidate names + features per CV.
# The model learns to CHOOSE among candidates, not to read raw text.
# ----------------------------------------------------------------------------
_FONT_CACHE = {}


def _font_map(path):
    """First-page spans -> list of (normalized_text, size, y_top). PDF only, fitz only."""
    if path in _FONT_CACHE:
        return _FONT_CACHE[path]
    result = _font_map_uncached(path)
    _FONT_CACHE[path] = result
    return result


def _font_map_uncached(path):
    if not path.lower().endswith(".pdf"):
        return [], 0.0
    try:
        import fitz
    except Exception:
        return [], 0.0
    spans = []
    try:
        d = fitz.open(path)
        page = d[0]
        page_h = page.rect.height or 1.0
        data = page.get_text("dict")
        for block in data.get("blocks", []):
            for line in block.get("lines", []):
                for s in line.get("spans", []):
                    t = normalize(s.get("text", ""))
                    if t:
                        spans.append((t, float(s.get("size", 0)), float(s.get("bbox", [0, 0, 0, 0])[1])))
        d.close()
        return spans, page_h
    except Exception:
        return [], 0.0


_HEADER_ALL = _HEADER_WORDS | {
    "experience", "education", "skills", "summary", "objective", "projects",
    "work", "academic", "record", "key", "achievements", "certifications",
}


_POOL_CACHE = {}


def candidate_pool(text, filename, path, drop_filename=False, include_reconcile=False):
    """Return list of dicts: {name, features{...}} for the ranker to choose among.

    drop_filename=True simulates ZIP/UUID uploads where the filename carries no name.
    include_reconcile=True adds the filename<->email reconciled candidate.

    Memoized per (path, drop_filename, include_reconcile): compare/train rebuild the
    same pools many times, so this avoids re-opening PDFs and re-running name-match.
    """
    ck = (path, drop_filename, include_reconcile)
    if ck in _POOL_CACHE:
        return _POOL_CACHE[ck]
    result = _candidate_pool_uncached(text, filename, path, drop_filename, include_reconcile)
    _POOL_CACHE[ck] = result
    return result


def _candidate_pool_uncached(text, filename, path, drop_filename=False, include_reconcile=False):
    import statistics
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    spans, page_h = _font_map(path)
    sizes = [s for _, s, _ in spans]
    med_size = statistics.median(sizes) if sizes else 0.0

    def font_feats(name):
        nn = normalize(name)
        if not nn or not spans:
            return 1.0, 0.5
        best = None
        for t, sz, y in spans:
            if nn in t or t in nn or sorted(nn.split()) == sorted(t.split()):
                if best is None or sz > best[0]:
                    best = (sz, y)
        if best is None:
            return 1.0, 0.5
        rel = (best[0] / med_size) if med_size else 1.0
        y_rel = (best[1] / page_h) if page_h else 0.5
        return rel, y_rel

    raw = []  # (name, source, line_idx)
    fn = m_filename(text, filename)
    if fn and not drop_filename:
        raw.append((fn, "filename", -1))
    if include_reconcile and not drop_filename:
        rec = m_reconcile(text, filename)
        if rec:
            raw.append((rec, "reconcile", -1))
    em = m_email(text, filename)
    if em:
        raw.append((em, "email", -1))
    if USE_SPACY:  # spaCy is a weak signal AND leaks memory over many calls; off by default
        try:
            nlp = _get_nlp()
            for ent in nlp(text[:600]).ents:
                if ent.label_ == "PERSON":
                    raw.append((ent.text.strip(), "spacy", 0))
        except Exception:
            pass
    for idx, ln in enumerate(lines[:10]):
        low = ln.lower()
        if "@" in ln or "http" in low or re.search(r"\d", ln):
            continue
        toks = ln.split()
        if 2 <= len(toks) <= 4 and all(re.fullmatch(r"[A-Za-z.]+", t) for t in toks):
            raw.append((" ".join(t.capitalize().rstrip(".") for t in toks), "topline", idx))

    # Deduplicate by normalized name, keep first occurrence's metadata.
    seen, cands = {}, []
    for name, src, idx in raw:
        key = normalize(name)
        if not key or key in seen:
            continue
        seen[key] = True
        cands.append({"name": name, "source": src, "line_idx": idx})

    # Per-candidate features (incl. agreement with the other candidates).
    for c in cands:
        nm = c["name"]
        nn = normalize(nm)
        toks = nn.split()
        agree = sum(1 for o in cands if o is not c and is_match(nm, o["name"]))
        rel, y_rel = font_feats(nm)
        c["features"] = {
            "src_filename": int(c["source"] == "filename"),
            "src_reconcile": int(c["source"] == "reconcile"),
            "src_email": int(c["source"] == "email"),
            "src_spacy": int(c["source"] == "spacy"),
            "src_topline": int(c["source"] == "topline"),
            "line_idx": float(c["line_idx"]),
            "n_tokens": float(len(toks)),
            "is_all_caps": int(nm.isupper()),
            "is_title": int(nm == nm.title()),
            "looks_header": int(any(w in nn.split() for w in _HEADER_ALL)),
            "agree_count": float(agree),
            "font_rel": float(rel),
            "y_rel": float(y_rel),
        }
    return cands


FEATURE_ORDER = [
    "src_filename", "src_reconcile", "src_email", "src_spacy", "src_topline",
    "line_idx", "n_tokens", "is_all_caps", "is_title", "looks_header",
    "agree_count", "font_rel", "y_rel",
]


def ensemble_pick(cands):
    """Heuristic choice over a candidate pool (mirrors m_ensemble, pool-based)."""
    if not cands:
        return None
    agreed = [c for c in cands if c["features"]["agree_count"] > 0]
    if agreed:
        return max(agreed, key=lambda c: c["features"]["agree_count"])["name"]
    for c in cands:
        if c["features"]["n_tokens"] == 2 and c["features"]["is_title"]:
            return c["name"]
    return cands[0]["name"]


# ----------------------------------------------------------------------------
# File discovery
# ----------------------------------------------------------------------------
def find_cvs(root):
    out = []
    for dirpath, _, files in os.walk(root):
        for f in files:
            if f.lower().endswith((".pdf", ".docx")):
                out.append(os.path.join(dirpath, f))
    return sorted(out)


# ----------------------------------------------------------------------------
# train mode — cross-validated ML ranker vs heuristic ensemble (offline only)
# ----------------------------------------------------------------------------
def run_train(args):
    import numpy as np
    from sklearn.linear_model import LogisticRegression
    from sklearn.ensemble import GradientBoostingClassifier
    from sklearn.model_selection import GroupKFold

    labels = {}
    with open(args.labels, encoding="utf-8") as f:
        for row in csv.DictReader(f):
            labels[row["filename"]] = row["true_name"]
    by_path = {os.path.basename(p): p for p in find_cvs(args.dir)}

    # Build candidate dataset.
    X, y, groups, meta = [], [], [], []
    per_cv = {}            # cv -> list of (cand_dict, is_correct)
    ceiling = 0            # CVs where at least one candidate is correct
    ensemble_correct = 0
    n = 0
    for fname, truth in labels.items():
        path = by_path.get(fname)
        if not path:
            continue
        n += 1
        text = extract_text(path)
        cands = candidate_pool(text, fname, path, drop_filename=args.no_filename)
        rows = []
        any_correct = False
        for c in cands:
            correct = is_match(c["name"], truth)
            any_correct = any_correct or correct
            X.append([c["features"][k] for k in FEATURE_ORDER])
            y.append(int(correct))
            groups.append(fname)
            meta.append((fname, c["name"]))
            rows.append((c, correct))
        per_cv[fname] = rows
        if any_correct:
            ceiling += 1
        # heuristic ensemble baseline over the same candidate pool
        if is_match(ensemble_pick(cands) or "", truth):
            ensemble_correct += 1

    X = np.array(X, dtype=float)
    y = np.array(y, dtype=int)
    print(f"CVs: {n} | candidate rows: {len(y)} | positives: {int(y.sum())}")
    print(f"Ceiling (>=1 correct candidate exists): {ceiling}/{n} ({ceiling/n:.0%})")
    print(f"Heuristic ensemble baseline           : {ensemble_correct}/{n} "
          f"({ensemble_correct/n:.0%})\n")

    models = {
        "logreg": LogisticRegression(max_iter=1000, class_weight="balanced"),
        "gboost": GradientBoostingClassifier(random_state=0),
    }
    n_splits = min(5, n)
    gkf = GroupKFold(n_splits=n_splits)
    uniq_groups = list(dict.fromkeys(groups))
    groups_arr = np.array(groups)

    for mname, model in models.items():
        cv_correct = 0
        for tr_idx, te_idx in gkf.split(X, y, groups_arr):
            model.fit(X[tr_idx], y[tr_idx])
            te_groups = set(groups_arr[te_idx])
            # score each test CV's candidates, pick highest
            for cv in te_groups:
                rows = per_cv[cv]
                if not rows:
                    continue
                feats = np.array([[c["features"][k] for k in FEATURE_ORDER]
                                  for c, _ in rows], dtype=float)
                try:
                    scores = model.predict_proba(feats)[:, 1]
                except Exception:
                    scores = model.decision_function(feats)
                best = int(np.argmax(scores))
                if rows[best][1]:        # chosen candidate is correct
                    cv_correct += 1
        print(f"ranker[{mname}] CV accuracy (group {n_splits}-fold): "
              f"{cv_correct}/{n} ({cv_correct/n:.0%})")

    # Feature importance from a full-data logreg fit (rough signal).
    lr = LogisticRegression(max_iter=1000, class_weight="balanced").fit(X, y)
    print("\nlogreg feature weights:")
    for k, w in sorted(zip(FEATURE_ORDER, lr.coef_[0]), key=lambda t: -abs(t[1])):
        print(f"  {k:14} {w:+.2f}")


# ----------------------------------------------------------------------------
# compare mode — (1) heuristics no-ML  (2) ML only  (3) both, head to head
# ----------------------------------------------------------------------------
def run_compare(args):
    import numpy as np
    from sklearn.ensemble import GradientBoostingClassifier
    from sklearn.linear_model import LogisticRegression
    from sklearn.model_selection import GroupKFold

    labels = {}
    with open(args.labels, encoding="utf-8") as f:
        for row in csv.DictReader(f):
            labels[row["filename"]] = row["true_name"]
    by_path = {os.path.basename(p): p for p in find_cvs(args.dir)}

    data = []
    for fname, truth in labels.items():
        path = by_path.get(fname)
        if path:
            data.append((fname, truth, extract_text(path), path))
    n = len(data)
    drop = args.no_filename

    # Config 1: heuristic, NO ML  (reconcile -> ensemble fallback)
    c1 = 0
    for fname, truth, text, path in data:
        rec = None if drop else m_reconcile(text, fname)
        if rec and is_match(rec, truth):
            c1 += 1
            continue
        pool = candidate_pool(text, fname, path, drop_filename=drop)
        if is_match(ensemble_pick(pool) or "", truth):
            c1 += 1

    feats_nofont = [k for k in FEATURE_ORDER if k not in ("font_rel", "y_rel")]

    def cv_rank(include_reconcile, factory, features=FEATURE_ORDER):
        per, X, y, g = {}, [], [], []
        for fname, truth, text, path in data:
            pool = candidate_pool(text, fname, path, drop_filename=drop,
                                  include_reconcile=include_reconcile)
            rows = []
            for c in pool:
                corr = is_match(c["name"], truth)
                X.append([c["features"][k] for k in features])
                y.append(int(corr)); g.append(fname); rows.append((c, corr))
            per[fname] = rows
        X = np.array(X, float); y = np.array(y, int); g = np.array(g)
        gkf = GroupKFold(n_splits=min(5, n))
        correct = 0
        for tr, te in gkf.split(X, y, g):
            m = factory(); m.fit(X[tr], y[tr])
            for cv in set(g[te]):
                rows = per[cv]
                if not rows:
                    continue
                f = np.array([[c["features"][k] for k in features]
                              for c, _ in rows], float)
                s = m.predict_proba(f)[:, 1]
                if rows[int(np.argmax(s))][1]:
                    correct += 1
        return correct

    # Font-only baseline: pick the candidate with the largest relative font.
    c_font = 0
    for fname, truth, text, path in data:
        pool = candidate_pool(text, fname, path, drop_filename=drop,
                              include_reconcile=True)
        if not pool:
            continue
        best = max(pool, key=lambda c: c["features"]["font_rel"])
        if is_match(best["name"], truth):
            c_font += 1

    gb = lambda: GradientBoostingClassifier(random_state=0)
    lr = lambda: LogisticRegression(max_iter=1000, class_weight="balanced")
    c2_nofont = cv_rank(True, gb, feats_nofont)      # ML, reconcile in pool, NO font
    c2_gb, c2_lr = cv_rank(False, gb), cv_rank(False, lr)
    c3_gb, c3_lr = cv_rank(True, gb), cv_rank(True, lr)

    def pct(x):
        return f"{x:>3}/{n}  {x/n:>4.0%}"
    mode = "NO filename (ZIP/UUID sim)" if drop else "with filename"
    print(f"\nDataset: {n} CVs   |   mode: {mode}")
    print("=" * 60)
    print(f"(1) Heuristics only (reconcile+ensemble), NO ML : {pct(c1)}")
    print(f"(F) Font-only pick (largest text)               : {pct(c_font)}")
    print(f"(2) ML only            gboost : {pct(c2_gb)}    logreg : {pct(c2_lr)}")
    print(f"(3) ML + reconcile     gboost : {pct(c3_gb)}    logreg : {pct(c3_lr)}")
    print(f"    ablation: ML+reconcile WITHOUT font (gboost): {pct(c2_nofont)}")
    print("=" * 60)


# ----------------------------------------------------------------------------
# autolabel mode — derive ground truth from filenames (then human verifies)
# ----------------------------------------------------------------------------
_MONTHS = {"jan", "feb", "mar", "march", "apr", "april", "may", "jun", "june",
           "jul", "july", "aug", "sep", "sept", "oct", "nov", "dec",
           "january", "february", "august", "september", "october",
           "november", "december"}
_LABEL_STOPWORDS = _FILENAME_NOISE | _MONTHS | {
    "batch", "current", "currently", "appraisal", "due", "ask", "notice",
    "period", "ctc", "lpa", "at", "is", "yrs", "year", "years", "exp",
    "experience", "fixed", "variable", "expected", "offer", "ok",
}


def _parse_name_segment(base):
    base = re.split(r"[\[\(]", base)[0]              # drop [10y_0m], (1), etc.
    base = re.sub(r"[_\-\.]+", " ", base)
    tokens = []                                      # camelCase-split each token
    for t in base.split():
        tokens += _camel(t) or [t]
    out = []
    for t in tokens:
        if re.search(r"\d", t):          # hit a number (batch/CTC) -> stop
            break
        low = t.lower()
        if low in _LABEL_STOPWORDS or low in _CERTS:  # noise / cert word
            if out:                       # noise after the name -> stop
                break
            continue                      # leading noise (Resume_, CV_) -> skip
        if not t.isalpha():
            break
        out.append(t.capitalize())
        if len(out) == 3:
            break
    return " ".join(out)


def derive_name_from_filename(filename):
    base = os.path.splitext(os.path.basename(filename))[0]
    if " - " in base:                               # "Orig - Candidate Name" export
        tail = _parse_name_segment(base.rsplit(" - ", 1)[-1])
        if tail:                                    # tail is the typed name...
            return tail
        return _parse_name_segment(base.rsplit(" - ", 1)[0])  # ...else the head
    return _parse_name_segment(base)


def run_autolabel(args):
    cvs = find_cvs(args.dir)
    rows = []
    for path in cvs:
        fname = os.path.basename(path)
        rows.append({"filename": fname, "true_name": derive_name_from_filename(fname)})
    with open(args.out, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=["filename", "true_name"])
        w.writeheader()
        w.writerows(rows)
    print(f"Wrote {len(rows)} auto-labels to {args.out}")
    print("VERIFY these before trusting eval results:\n")
    for r in rows:
        print(f"  {r['true_name']:<28} <- {r['filename']}")


# ----------------------------------------------------------------------------
# label mode
# ----------------------------------------------------------------------------
def run_label(args):
    cvs = find_cvs(args.dir)
    if not cvs:
        print(f"No PDF/DOCX files under {args.dir}")
        return
    print(f"Labeling {len(cvs)} CV(s). Press Enter to accept a suggestion, "
          f"type a correction, or 's' to skip.\n")
    rows = []
    for i, path in enumerate(cvs, 1):
        text = extract_text(path)
        fname = os.path.basename(path)
        print("=" * 70)
        print(f"[{i}/{len(cvs)}] {fname}")
        for ln in [l for l in text.splitlines() if l.strip()][:15]:
            print("   |", ln[:90])
        suggestion = ""
        if args.seed == "filename":
            suggestion = m_filename(text, fname) or ""
        elif args.seed == "gemini":
            suggestion = m_gemini(text, fname) or ""
        prompt = f"  TRUE NAME [{suggestion}]: " if suggestion else "  TRUE NAME: "
        ans = input(prompt).strip()
        if ans.lower() == "s":
            continue
        true_name = ans or suggestion
        if not true_name:
            print("  (no name entered, skipping)")
            continue
        rows.append({"filename": fname, "true_name": true_name})
    with open(args.out, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=["filename", "true_name"])
        w.writeheader()
        w.writerows(rows)
    print(f"\nWrote {len(rows)} labels to {args.out}")


# ----------------------------------------------------------------------------
# eval mode
# ----------------------------------------------------------------------------
def run_eval(args):
    labels = {}
    with open(args.labels, encoding="utf-8") as f:
        for row in csv.DictReader(f):
            labels[row["filename"]] = row["true_name"]
    if not labels:
        print(f"No labels found in {args.labels}")
        return

    methods = list(LOCAL_METHODS)
    if not args.no_gemini:
        methods.append(GEMINI_METHOD)

    print(f"indian-namematch available: {HAS_INM}")
    print(f"Evaluating {len(labels)} labeled CV(s) across "
          f"{len(methods)} method(s)...\n")

    by_path = {os.path.basename(p): p for p in find_cvs(args.dir)}
    correct = {name: 0 for name, _ in methods}
    empty = {name: 0 for name, _ in methods}
    any_local_correct = 0
    gemini_only = 0
    nobody = 0
    detail_rows = []
    n = 0

    for fname, truth in labels.items():
        path = by_path.get(fname)
        if not path:
            print(f"  [skip] file not found for label: {fname}")
            continue
        n += 1
        text = extract_text(path)
        row = {"filename": fname, "true_name": truth}
        local_hit = False
        gem_hit = False
        for mname, fn in methods:
            pred = fn(text, fname) or ""
            hit = is_match(pred, truth)
            row[f"{mname}_pred"] = pred
            row[f"{mname}_match"] = "Y" if hit else ""
            if not pred:
                empty[mname] += 1
            if hit:
                correct[mname] += 1
                if mname == "gemini":
                    gem_hit = True
                else:
                    local_hit = True
        if local_hit:
            any_local_correct += 1
        elif gem_hit:
            gemini_only += 1
        else:
            nobody += 1
        detail_rows.append(row)

    if n == 0:
        print("No labeled files matched files in --dir.")
        return

    # Per-CV detail CSV
    fieldnames = ["filename", "true_name"]
    for mname, _ in methods:
        fieldnames += [f"{mname}_pred", f"{mname}_match"]
    with open(args.out, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames)
        w.writeheader()
        w.writerows(detail_rows)

    # Summary
    print("=" * 60)
    print(f"{'METHOD':<12}{'ACCURACY':>12}{'EMPTY/NO-NAME':>16}")
    print("-" * 60)
    for mname, _ in methods:
        acc = correct[mname] / n
        print(f"{mname:<12}{correct[mname]:>4}/{n:<4}{acc:>6.0%}   "
              f"{empty[mname]:>8}/{n}")
    print("=" * 60)
    print(f"Total labeled CVs evaluated : {n}")
    print(f"Solved by a LOCAL method    : {any_local_correct:>4}  ({any_local_correct/n:.0%})")
    if not args.no_gemini:
        print(f"Only Gemini got it right    : {gemini_only:>4}  ({gemini_only/n:.0%})  <- residual quota need")
    print(f"No method got it right      : {nobody:>4}  ({nobody/n:.0%})  <- needs manual review")
    print("=" * 60)
    print(f"Per-CV details written to {args.out}")


# ----------------------------------------------------------------------------
def main():
    p = argparse.ArgumentParser(description="QC harness for CV name extraction.")
    sub = p.add_subparsers(dest="cmd", required=True)

    pa = sub.add_parser("autolabel", help="Derive labels from filenames (verify after).")
    pa.add_argument("--dir", required=True)
    pa.add_argument("--out", default="labels.csv")
    pa.set_defaults(func=run_autolabel)

    pt = sub.add_parser("train", help="Cross-validated ML ranker vs heuristic ensemble.")
    pt.add_argument("--dir", required=True)
    pt.add_argument("--labels", default="labels.csv")
    pt.add_argument("--pdf-engine", choices=["auto", "fitz", "pdfminer", "pypdfium2"],
                    default="pdfminer")
    pt.add_argument("--no-filename", action="store_true",
                    help="Drop filename candidate (simulate ZIP/UUID uploads).")
    pt.set_defaults(func=run_train)

    pc = sub.add_parser("compare", help="Heuristics-only vs ML-only vs both.")
    pc.add_argument("--dir", required=True)
    pc.add_argument("--labels", default="labels.csv")
    pc.add_argument("--pdf-engine", choices=["auto", "fitz", "pdfminer", "pypdfium2"],
                    default="pdfminer")
    pc.add_argument("--no-filename", action="store_true")
    pc.set_defaults(func=run_compare)

    pl = sub.add_parser("label", help="Create ground-truth labels for a folder of CVs.")
    pl.add_argument("--dir", required=True)
    pl.add_argument("--out", default="labels.csv")
    pl.add_argument("--seed", choices=["none", "filename", "gemini"], default="filename")
    pl.set_defaults(func=run_label)

    pe = sub.add_parser("eval", help="Score every method against the labels.")
    pe.add_argument("--dir", required=True)
    pe.add_argument("--labels", default="labels.csv")
    pe.add_argument("--out", default="results.csv")
    pe.add_argument("--no-gemini", action="store_true")
    pe.add_argument("--pdf-engine", choices=["auto", "fitz", "pdfminer", "pypdfium2"],
                    default="auto")
    pe.set_defaults(func=run_eval)

    args = p.parse_args()
    if getattr(args, "pdf_engine", None):
        globals()["PDF_ENGINE"] = args.pdf_engine
    args.func(args)


if __name__ == "__main__":
    main()
